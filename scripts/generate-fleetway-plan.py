from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"C:\desenvolvimento\rqcode\projetos\fleetway\PLANO-EXECUCAO-FLEETWAY.docx")
BLUE = "123B5D"
TEAL = "00A6A6"
LIGHT = "E8EEF5"
RED = "9B1C1C"
GRAY = "5B6573"

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)

def set_cell_text(cell, text, bold=False, color="000000", size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.55)
sec.left_margin = sec.right_margin = Inches(0.65)
sec.header_distance = sec.footer_distance = Inches(0.3)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(9.5)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.08
for name, size, before, after in (("Heading 1", 15, 10, 5), ("Heading 2", 11.5, 7, 3)):
    s = styles[name]
    s.font.name = "Aptos Display"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(BLUE)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)

header = sec.header.paragraphs[0]
header.text = "RQCODE  |  FLEETWAY"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs:
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(2)
r = title.add_run("FLEETWAY")
r.font.name = "Aptos Display"; r.font.size = Pt(25); r.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE)
sub = doc.add_paragraph("PLANO DE EXECUÇÃO E ACOMPANHAMENTO  •  30/07/2026")
sub.paragraph_format.space_after = Pt(8)
for r in sub.runs:
    r.font.size = Pt(9); r.bold = True; r.font.color.rgb = RGBColor.from_string(TEAL)

t = doc.add_table(rows=1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
widths = (2.35, 2.35, 2.35)
for cell, width in zip(t.rows[0].cells, widths): cell.width = Inches(width)
cards = (("OBJETIVO", "Lançar um SaaS comercial, seguro e administrado pelo RQCODE."),
         ("REGRA DO DIA", "Escolher 3 entregas. Concluir antes de abrir novas frentes."),
         ("FONTE DA VERDADE", "GitHub + RQCODE. Nenhum segredo dentro do repositório."))
for cell, (label, value) in zip(t.rows[0].cells, cards):
    shade(cell, LIGHT); set_cell_text(cell, f"{label}\n{value}", bold=False, color=BLUE, size=9)

doc.add_heading("Situação verificada", level=1)
facts = [
    ("VPS", "SSH por chave OK; Nginx e PHP-FPM ativos; disco em 40%."),
    ("Produto", "Multi-tenant, 4 painéis, i18n, voz, pagamentos e suporte já estruturados."),
    ("Segurança", "54 testes passaram; banco local indisponível bloqueia validação completa."),
    ("Git", "Fleetway local e produção ainda não são repositórios Git."),
    ("Crítico", "Credenciais legadas em texto claro: rotacionar e retirar do fluxo operacional."),
]
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
for cell, width in zip(table.rows[0].cells, (1.25, 5.8)): cell.width = Inches(width)
for i, h in enumerate(("ÁREA", "DIAGNÓSTICO")): shade(table.rows[0].cells[i], BLUE); set_cell_text(table.rows[0].cells[i], h, True, "FFFFFF")
for label, value in facts:
    cells = table.add_row().cells
    cells[0].width, cells[1].width = Inches(1.25), Inches(5.8)
    set_cell_text(cells[0], label, True, BLUE)
    set_cell_text(cells[1], value, color=RED if label == "Crítico" else "000000")

doc.add_heading("Sequência obrigatória", level=1)
phases = [
    ("1", "SEGURANÇA", "Rotacionar segredos • cofre • MFA • APP_DEBUG=false • backup/restauração"),
    ("2", "GITHUB", "Criar repo privado • preparar commit • revisar segredos • proteger main • CI"),
    ("3", "PRODUTO", "Onboarding • Asaas • suporte IA • PT/EN/ES • voz segura • UX responsiva"),
    ("4", "RQCODE", "API/webhooks • contas • usuários • planos • cobrança • suporte • auditoria"),
    ("5", "LANÇAMENTO", "Validar • backup • deploy • smoke • monitorar • comercializar"),
]
table = doc.add_table(rows=1, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
for cell, w in zip(table.rows[0].cells, (0.5, 1.25, 5.3)): cell.width = Inches(w)
for i, h in enumerate(("#", "FASE", "ENTREGAS")): shade(table.rows[0].cells[i], TEAL); set_cell_text(table.rows[0].cells[i], h, True, "FFFFFF")
for n, phase, work in phases:
    cells = table.add_row().cells
    cells[0].width, cells[1].width, cells[2].width = Inches(0.5), Inches(1.25), Inches(5.3)
    set_cell_text(cells[0], n, True, TEAL, 10); set_cell_text(cells[1], phase, True, BLUE); set_cell_text(cells[2], work)

doc.add_page_break()
doc.add_heading("Checklist de mesa", level=1)
sections = {
    "A. Segurança e acessos": ["Rotacionar credenciais legadas", "Ativar MFA e cofre", "Testar restauração de backup", "Revisar usuários root/admin"],
    "B. Fonte da verdade": ["Criar GitHub privado fleetway", "Preparar commit com publish-github.bat", "Revisar e fazer push autorizado", "Configurar CI e proteção da main"],
    "C. Produto comercial": ["Asaas ponta a ponta", "Suporte IA + humano", "PT-BR, EN e ES completos", "Voz com confirmação", "Mobile, acessibilidade e UX aprovados"],
    "D. RQCODE central": ["Sincronizar contas e usuários", "Centralizar planos e cobranças", "Centralizar suporte/SLA", "Exibir MRR, churn e saúde"],
    "E. Operação": ["Banco local ativo", "validate.bat aprovado", "deploy-vps.bat com backup", "Smoke e monitoramento aprovados"],
}
for heading, items in sections.items():
    doc.add_heading(heading, level=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    for idx, item in enumerate(items):
        run = p.add_run(f"☐ {item}" + ("     " if idx % 2 == 0 else "\n"))
        run.font.name = "Aptos"; run.font.size = Pt(9.5)

doc.add_heading("Comandos", level=1)
cmds = [
    ("Iniciar local", r"C:\desenvolvimento\fleetway\start-local.bat"),
    ("Validar", r"C:\desenvolvimento\fleetway\validate.bat"),
    ("Preparar Git", r"publish-github.bat -RepositoryUrl <URL_PRIVADA>"),
    ("Deploy", r"C:\desenvolvimento\fleetway\deploy-vps.bat"),
]
table = doc.add_table(rows=1, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
for cell, width in zip(table.rows[0].cells, (1.55, 5.5)): cell.width = Inches(width)
for i, h in enumerate(("AÇÃO", "COMANDO")): shade(table.rows[0].cells[i], BLUE); set_cell_text(table.rows[0].cells[i], h, True, "FFFFFF")
for action, cmd in cmds:
    cells = table.add_row().cells; cells[0].width, cells[1].width = Inches(1.55), Inches(5.5); set_cell_text(cells[0], action, True, BLUE); set_cell_text(cells[1], cmd, size=8.5)

doc.add_heading("Ritual profissional", level=1)
p = doc.add_paragraph("SEGUNDA: escolher marco e métrica.  •  DIARIAMENTE: 3 entregas e 1 registro.  •  SEXTA: demonstrar, medir e decidir.\nNUNCA: publicar sem teste/backup; guardar senha em arquivo; abrir nova frente antes de concluir o marco atual.")
p.paragraph_format.space_after = Pt(4)

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("RQCODE • Fleetway • imprimir e manter visível")
r.font.name = "Aptos"; r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(GRAY)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
